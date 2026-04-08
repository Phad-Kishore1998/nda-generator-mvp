"""
exporter.py — Optional export of the final NDA document to .docx format.
"""

from pathlib import Path

FINAL_DIR = Path(__file__).parent.parent.parent / "outputs" / "final"
FINAL_FILE_TXT = FINAL_DIR / "nda_final.txt"
FINAL_FILE_DOCX = FINAL_DIR / "nda_final.docx"


def export_to_docx(txt_path: str = None) -> str:
    """
    Convert the final NDA .txt file to a formatted .docx document.

    Args:
        txt_path: Optional override path to the .txt file.

    Returns:
        Path to the generated .docx file.

    Raises:
        FileNotFoundError: If the source .txt file does not exist.
        ImportError: If python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    source = Path(txt_path) if txt_path else FINAL_FILE_TXT
    if not source.exists():
        raise FileNotFoundError(f"Source file not found: {source}")

    with open(source, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("=" * 40):
            i += 1
            continue

        if line.startswith("-" * 40):
            doc.add_paragraph()
            i += 1
            continue

        if line == "NON-DISCLOSURE AGREEMENT":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(18)
            i += 1
            continue

        if line.startswith("Generated on:") or line.startswith("NOTICE:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line).italic = True
            i += 1
            continue

        if line == "SIGNATURE PAGE":
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue

        if line == "END OF AGREEMENT":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line).bold = True
            i += 1
            continue

        # Detect section headings (lines with numbered clause pattern like "1." at start
        # or all-caps short lines that look like headings)
        if line and len(line) < 80 and (
            line.isupper() or
            (line[0].isdigit() and "." in line[:3] and len(line.split(".")[0]) <= 2)
        ):
            p = doc.add_heading(line, level=2)
            i += 1
            continue

        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()

        i += 1

    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(FINAL_FILE_DOCX))

    print(f"📄 DOCX export saved → {FINAL_FILE_DOCX}")
    return str(FINAL_FILE_DOCX)
